from pathlib import Path
import html, re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from ebooklib import epub

ROOT = Path(__file__).resolve().parents[1]
FINAL = ROOT / 'balance-due' / 'final'
SRC = FINAL / 'Balance_Due_FINAL.txt'
DOCX = FINAL / 'Balance_Due_FINAL.docx'
EPUB = FINAL / 'Balance_Due_FINAL.epub'
text = SRC.read_text(encoding='utf-8').replace('\r\n','\n')

# The canonical TXT begins with title/author/copyright front matter.
body_marker = '\nPrologue\n'
pos = text.find(body_marker)
if pos < 0:
    raise SystemExit('Could not find Prologue in canonical manuscript')
front = text[:pos].strip()
body = text[pos+1:].strip()

major = {'Prologue', 'Epilogue'} | {f'Chapter {n}' for n in range(1,10)}
subs = {'The Historical Record', 'The Double Event'}

# ---------- DOCX ----------
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(6), Inches(9)
sec.top_margin = sec.bottom_margin = Inches(.70)
sec.left_margin = sec.right_margin = Inches(.72)
sec.footer_distance = Inches(.30)

normal = doc.styles['Normal']
normal.font.name = 'Garamond'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Garamond')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Garamond')
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.08
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.first_line_indent = Inches(.22)
normal.paragraph_format.widow_control = True

for sname, size, italic in [('Heading 1',17,False), ('Heading 2',12.5,True)]:
    st = doc.styles[sname]
    st.font.name='Garamond'; st.font.size=Pt(size); st.font.bold=(sname=='Heading 1'); st.font.italic=italic
    st._element.rPr.rFonts.set(qn('w:ascii'),'Garamond'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Garamond')
    st.paragraph_format.first_line_indent=Inches(0); st.paragraph_format.keep_with_next=True
    st.paragraph_format.space_after=Pt(14 if sname=='Heading 1' else 16)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(165); p.paragraph_format.first_line_indent=Inches(0)
r=p.add_run('BALANCE DUE'); r.bold=True; r.font.name='Garamond'; r.font.size=Pt(28)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18); p.paragraph_format.first_line_indent=Inches(0)
r=p.add_run('TOBY BRANDON'); r.font.name='Garamond'; r.font.size=Pt(13)

doc.add_page_break()
for _ in range(8): doc.add_paragraph('')
for line in [
    'Copyright © 2026 Toby Brandon', 'All rights reserved.', '',
    'This is a work of fiction. Names, characters, businesses, places, events, and incidents are either products of the author’s imagination or are used fictitiously. Any resemblance to actual persons, living or dead, is coincidental except where historical events or places are deliberately referenced as part of the fictional setting.',
    '', 'First revised edition, 2026.'
]:
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Inches(0); p.paragraph_format.space_after=Pt(7)
    if line: p.add_run(line)

main_sec=doc.add_section(WD_SECTION.NEW_PAGE)
main_sec.page_width, main_sec.page_height=Inches(6), Inches(9)
main_sec.top_margin=main_sec.bottom_margin=Inches(.70); main_sec.left_margin=main_sec.right_margin=Inches(.72); main_sec.footer_distance=Inches(.30)
main_sec.footer.is_linked_to_previous=False
pg=OxmlElement('w:pgNumType'); pg.set(qn('w:start'),'1'); main_sec._sectPr.append(pg)
fp=main_sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=fp.add_run(); a=OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'),'begin'); b=OxmlElement('w:instrText'); b.set(qn('xml:space'),'preserve'); b.text=' PAGE '; c=OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'),'end'); run._r.extend([a,b,c])

first_major=True; first_after=False; current=None
for raw in body.splitlines():
    line=raw.strip()
    if not line: continue
    if line in major:
        current=line
        if not first_major: doc.add_page_break()
        first_major=False
        p=doc.add_paragraph(style='Heading 1'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(line.upper()); first_after=True
    elif line in subs:
        p=doc.add_paragraph(style='Heading 2'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(line); first_after=True
    else:
        p=doc.add_paragraph(line, style='Normal')
        if current=='Chapter 5': p.paragraph_format.line_spacing=1.04
        if first_after: p.paragraph_format.first_line_indent=Inches(0); first_after=False

doc.core_properties.title='Balance Due'; doc.core_properties.author='Toby Brandon'; doc.core_properties.subject='Victorian psychological crime fiction'
doc.save(DOCX)

# ---------- EPUB ----------
book=epub.EpubBook(); book.set_identifier('balance-due-toby-brandon-2026'); book.set_title('Balance Due'); book.set_language('en-GB'); book.add_author('Toby Brandon')
css='''body{font-family:serif;line-height:1.35}p{text-indent:1.2em;margin:0 0 .35em}h1{text-align:center;margin-top:15%;margin-bottom:1.5em}h2{text-align:center;font-style:italic;font-weight:normal;margin-bottom:1.5em}.copyright p{text-indent:0;margin:.7em 0}'''
style=epub.EpubItem(uid='style', file_name='styles/book.css', media_type='text/css', content=css.encode()); book.add_item(style)

title=epub.EpubHtml(title='Balance Due', file_name='title.xhtml', lang='en-GB')
title.content='<html><body><h1>BALANCE DUE</h1><p style="text-align:center;text-indent:0">TOBY BRANDON</p></body></html>'; title.add_item(style); book.add_item(title)
copy=epub.EpubHtml(title='Copyright', file_name='copyright.xhtml', lang='en-GB')
copy.content='<html><body class="copyright"><p>Copyright © 2026 Toby Brandon</p><p>All rights reserved.</p><p>This is a work of fiction. Names, characters, businesses, places, events, and incidents are either products of the author’s imagination or are used fictitiously. Any resemblance to actual persons, living or dead, is coincidental except where historical events or places are deliberately referenced as part of the fictional setting.</p><p>First revised edition, 2026.</p></body></html>'; copy.add_item(style); book.add_item(copy)

sections=[]; cur_title=None; cur_sub=None; paras=[]
def flush():
    global cur_title,cur_sub,paras
    if not cur_title: return
    idx=len(sections)+1; fname=f'section-{idx:02d}.xhtml'; ch=epub.EpubHtml(title=(cur_title if not cur_sub else f'{cur_title} — {cur_sub}'), file_name=fname, lang='en-GB')
    pieces=[f'<h1>{html.escape(cur_title)}</h1>']
    if cur_sub: pieces.append(f'<h2>{html.escape(cur_sub)}</h2>')
    pieces.extend(f'<p>{html.escape(p)}</p>' for p in paras)
    ch.content='<html><body>'+''.join(pieces)+'</body></html>'; ch.add_item(style); book.add_item(ch); sections.append(ch)
    cur_title=cur_sub=None; paras=[]
for raw in body.splitlines():
    line=raw.strip()
    if not line: continue
    if line in major:
        flush(); cur_title=line
    elif line in subs and cur_title:
        cur_sub=line
    elif cur_title:
        paras.append(line)
flush()
book.toc=tuple(sections); book.spine=['nav', title, copy] + sections
book.add_item(epub.EpubNcx()); book.add_item(epub.EpubNav()); epub.write_epub(str(EPUB), book, {})

print(f'Built {DOCX.relative_to(ROOT)} and {EPUB.relative_to(ROOT)} from canonical TXT.')
